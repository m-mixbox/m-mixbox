from turtle import width
from docx.shared import Inches
from docx import Document
from flask import Flask, request, send_file, jsonify
import os
import random
import string
from io import BytesIO
from docx import Document
from docx.shared import Inches
from pathlib import Path
import traceback   

def abc():
        values_list = ["hello" for x in range(50)]
        template_path = os.path.join(os.path.dirname(__file__), "others.docx")
        doc = Document(template_path)

        table = doc.tables[0]
        table_input_position = [3, 3, 3, 3, 3, 3, 1, 3]

        for i in range(len(table.rows)):
            #row = table.rows[i]
            #row_data = [cell.text.strip() for cell in row.cells]
            #print(len(row_data))
            #print(row_data)
            if i != 0:
                print(table.cell(i, table_input_position[i-1]).text)
                table.cell(i, table_input_position[i-1]).text = values_list[i]
            #elif i==5 :
                #table.cell(i, 5).text = ""
            #elif i==6 :
                #table.cell(i, 5).text = ""
            #elif i == 7:
                #table.cell(i, 5).text = values_list[5]
            #elif i ==8 :
                #table.cell(i, 5).text = values_list[6]
            #elif i==23:
                #table.cell(i, 5).text = ""

        #print("+++++++++++++++++++++++")
        #table_input_position_2 = [3,3,3,3,3,3,3,3,3,3,3,3,3,3,3,3,3]
        table_2 = doc.tables[1]
        #table_input_position_2 = [2, 2, 2]
        for i in range(len(table_2.rows)):
            #print(i)
            #row = table_2.rows[i]
            #row_data = [cell.text.strip() for cell in row.cells]
            #print(len(row_data))
            #print(row_data)
            if 1<= i <= 5:
                table_2.cell(i, 2).text = values_list[i+9]
                
            #elif i == 16:
                #table_2.cell(i, 2).text = ""
            #elif i == 18:
                #table_2.cell(i, 2).text = ""
        #print("+++++++++++++++++++++++")
        #table_input_position_3 = [4,4,4,4,4,4,4,4,4]
        table_3 = doc.tables[2]
        #table_input_position_2 = [2, 2, 2]
        for i in range(len(table_3.rows)):
            #row = table_3.rows[i]
            #row_data = [cell.text.strip() for cell in row.cells]
            #print(len(row_data))
            #print(row_data)
            if i !=0 or i != 5:
                table_3.cell(i, 3).text = values_list[i+25]
        
        table_4 = doc.tables[3]
        #table_input_position_2 = [2, 2, 2]
        for i in range(len(table_4.rows)):
            #row = table_4.rows[i]
            #row_data = [cell.text.strip() for cell in row.cells]
            #print(len(row_data))
            #print(row_data)
            table_3.cell(i, 0).text = values_list[i+36]

        #table_4 = doc.tables[3]
        #for row in table_4.rows:
            #for cell in row.cells:
                #for paragraph in cell.paragraphs:
                    #for run in paragraph.runs:
                        #run.clear()
abc()