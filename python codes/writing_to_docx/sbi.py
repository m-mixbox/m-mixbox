from turtle import width
from docx.shared import Inches
from docx import Document
from flask import Flask, request, send_file, jsonify
import os

# Open the existing Word document
#local_path = os.path.dirname(__file__)
#image_dir = os.path.join(local_path,"images")
#output_dir = os.path.join(local_path,"generated")
#doc = Document(os.path.join(local_path,"SBI Format.docx"))

# Access the first table (index 0)
def sbi_format(local_path,image_dir,output_dir,values_list,random_filename,output_path):
        
        template_path = os.path.join(os.path.dirname(__file__), "SBI Format.docx")
        print(f"Template path: {template_path}")

        if not os.path.exists(template_path):
            return jsonify({"error": "Template file 'SBI Format.docx' not found."}), 404

        # Open the template
        doc = Document(template_path)

        table = doc.tables[0]
        table_input_position = [5, 5, 5, 5, 5, 5, 5, 5, 5]

        for i in range(len(table.rows)):
            if i <= 4:
                table.cell(i, table_input_position[i]).text = values_list[i]
            elif i==5 :
                table.cell(i, 5).text = ""
            elif i==6 :
                table.cell(i, 5).text = ""
            elif i == 7:
                table.cell(i, 5).text = values_list[5]
            elif i ==8 :
                table.cell(i, 5).text = values_list[6]
            elif i==23:
                table.cell(i, 5).text = ""


        table_2 = doc.tables[1]
        table_input_position_2 = [2, 2, 2]
        for i in range(len(table_2.rows)):
            if 12 <= i < 15:
                table_2.cell(i, table_input_position_2[i - 12]).text = ""
            elif i == 16:
                table_2.cell(i, 2).text = ""
            elif i == 18:
                table_2.cell(i, 2).text = ""

        table_3 = doc.tables[2]
        for row in table_3.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.clear()

        image_index = 0
        start_row = 0
        start_col = 0
        image_paths = [os.path.join(image_dir,"1E2KS.jpg"),os.path.join(image_dir,"1X0I8.jpg"),os.path.join(image_dir,"2A1VW.jpg"),os.path.join(image_dir,"2GW4H.jpg")]

        for row_idx in range(start_row, len(table_3.rows)):
            for col_idx in range(start_col, len(table_3.columns)):
                if image_index >= len(image_paths):
                    break
                cell = table_3.cell(row_idx, col_idx)
                paragraph = cell.paragraphs[0]
                run = paragraph.add_run()
                try:
                    run.add_picture(image_paths[image_index], width=Inches(4.06))
                except Exception as e:
                    print(f"Error adding image {image_paths[image_index]}: {e}")
                image_index += 1
            if image_index >= len(image_paths):
                break

        # Save the document to a file
        doc.save(output_path)
        return output_path

        print(f"File saved to {output_path}")
