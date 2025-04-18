from fastapi import FastAPI, UploadFile, Form
from fastapi.responses import FileResponse, JSONResponse
from fastapi.middleware.cors import CORSMiddleware
from typing import Dict
from pydantic import BaseModel
import os
from pathlib import Path
from docx import Document
from docx.shared import Inches

app = FastAPI()

# Enable CORS for development purposes
app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

# Directory to save and serve files
generated_dir = Path("generated")
generated_dir.mkdir(exist_ok=True)

# Directory to store images
images_dir = Path("images")
images_dir.mkdir(exist_ok=True)

class Payload(BaseModel):
    data: Dict[str, str]

@app.post("/generate-docx/")
async def generate_docx():
    try:
        # Save uploaded template
        template_path = os.path.join(os.path.dirname(__file__), "SBI Format.docx")

        # Open the template
        doc = Document(template_path)

        # Additional functionality: Modify tables and insert images
        # Access the first table (index 0)
        table = doc.tables[0]
        table_input_position = [5, 5, 5, 5, 5, 5, 5, 5, 5]

        # Add text to specific cells
        for i in range(len(table.rows)):
            if i <= 8:
                table.cell(i, table_input_position[i]).text = "  hello"
            elif i == 23:
                table.cell(i, 5).text = "  hello"

        # Access the second table (index 1)
        table_2 = doc.tables[1]
        table_input_position_2 = [2, 2, 2]
        for i in range(len(table_2.rows)):
            if 12 <= i < 15:
                table_2.cell(i, table_input_position_2[i - 12]).text = "  hello"
            elif i == 16:
                table_2.cell(i, 2).text = "  hello"
            elif i == 18:
                table_2.cell(i, 2).text = "  hello"

        # Access the third table (index 2)
        table_3 = doc.tables[2]
        for row in table_3.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.clear()

        # Add images to the third table
        image_index = 0  # Track which image to add
        start_row = 0
        start_col = 0
        image_paths = [
            str(images_dir / "1E2KS.jpg"),
            str(images_dir / "1X0I8.jpg"),
            str(images_dir / "2A1VW.jpg"),
            str(images_dir / "2GW4H.jpg")
        ]

        for row_idx in range(start_row, len(table_3.rows)):
            for col_idx in range(start_col, len(table_3.columns)):
                if image_index >= len(image_paths):
                    break  # Stop if no more images to add
                
                cell = table_3.cell(row_idx, col_idx)
                paragraph = cell.paragraphs[0]  # Get the first paragraph in the cell
                run = paragraph.add_run()
                run.add_picture(image_paths[image_index], width=Inches(4.06))  # Adjust width as needed
                image_index += 1  # Move to the next image
            if image_index >= len(image_paths):
                break

        # Save the generated document
        generated_filename = f"generated_file.docx"
        generated_path = generated_dir / generated_filename
        doc.save(generated_path)

        # Return download link
        return JSONResponse({"message": "File generated successfully", "download_url": f"/download/{generated_filename}"})

    except Exception as e:
        return JSONResponse({"error": str(e)}, status_code=500)

@app.get("/download/{filename}")
async def download_file(filename: str):
    file_path = generated_dir / filename
    if file_path.exists():
        return FileResponse(file_path, filename=filename)
    return JSONResponse({"error": "File not found"}, status_code=404)

# Run the app with: uvicorn script_name:app --reload
