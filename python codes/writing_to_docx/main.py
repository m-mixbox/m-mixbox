import os
import random
import string
from flask import Flask, request, send_file, jsonify
from io import BytesIO
from docx import Document
from docx.shared import Inches
from pathlib import Path
import traceback
from flask_cors import CORS

app = Flask(__name__)
CORS(app, resources={r"/*": {"origins": "*"}})

@app.before_request
def log_request_info():
    print(f"Request Method: {request.method}")
    print(f"Request Headers: {request.headers}")
    print(f"Request Origin: {request.headers.get('Origin')}")

@app.after_request
def after_request(response):
    print(response.headers)
    return response


# Directory to store images and generated files
images_dir = Path(os.path.dirname(__file__)) / "images"
output_dir = Path(os.path.dirname(__file__)) / "generated"

# Ensure the output directory exists
if not output_dir.exists():
    output_dir.mkdir(parents=True)

# Helper function to generate a random 7-character string for filenames
def generate_random_filename(length=7):
    return ''.join(random.choices(string.ascii_letters + string.digits, k=length))

@app.route('/api/data', methods=['OPTIONS'])
def handle_options():
    return '', 204  # HTTP 204 No Content

@app.route('/generate-docx/', methods=['POST'])
def generate_docx():
    try:
        print("Starting DOCX generation process...")

        # Load the template DOCX file
        template_path = Path(os.path.dirname(__file__)) / "SBI Format.docx"
        print(f"Template path: {template_path}")

        if not template_path.exists():
            return jsonify({"error": "Template file 'SBI Format.docx' not found."}), 404

        # Open the template
        doc = Document(template_path)
        print("Started docx editing")

        # Modify the tables with specific content (same as original code)
        table = doc.tables[0]
        table_input_position = [5, 5, 5, 5, 5, 5, 5, 5, 5]

        for i in range(len(table.rows)):
            if i <= 8:
                table.cell(i, table_input_position[i]).text = "  hello"
            elif i == 23:
                table.cell(i, 5).text = "  hello"

        table_2 = doc.tables[1]
        table_input_position_2 = [2, 2, 2]
        for i in range(len(table_2.rows)):
            if 12 <= i < 15:
                table_2.cell(i, table_input_position_2[i - 12]).text = "  hello"
            elif i == 16:
                table_2.cell(i, 2).text = "  hello"
            elif i == 18:
                table_2.cell(i, 2).text = "  hello"

        table_3 = doc.tables[2]
        for row in table_3.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.clear()
        print("Edited the docx file")

        # List of image paths to add to the document
        image_index = 0
        start_row = 0
        start_col = 0
        image_paths = [
            Path(images_dir) / "1E2KS.jpg",
            Path(images_dir) / "1X0I8.jpg",
            Path(images_dir) / "2A1VW.jpg",
            Path(images_dir) / "2GW4H.jpg"
        ]

        # Check if image files exist
        for image_path in image_paths:
            if not image_path.exists():
                print(f"Warning: Image file not found: {image_path}")

        # Add images to the docx file
        for row_idx in range(start_row, len(table_3.rows)):
            for col_idx in range(start_col, len(table_3.columns)):
                if image_index >= len(image_paths):
                    break
                cell = table_3.cell(row_idx, col_idx)
                paragraph = cell.paragraphs[0]
                run = paragraph.add_run()
                try:
                    run.add_picture(str(image_paths[image_index]), width=Inches(4.06))
                except Exception as e:
                    print(f"Error adding image {image_paths[image_index]}: {e}")
                image_index += 1
            if image_index >= len(image_paths):
                break
        print("Added images to the docx file")

        # Generate a random filename for the output file
        random_filename = generate_random_filename() + ".docx"
        output_path = output_dir / random_filename

        # Save the document to a file
        doc.save(output_path)

        print(f"File saved to {output_path}")

        # Send the saved file as a download
        return send_file(
            output_path,
            as_attachment=True,
            download_name=random_filename,
            mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        )

    except Exception as e:
        print(f"Error occurred: {str(e)}")
        traceback.print_exc()
        return jsonify({"error": f"Internal Server Error: {str(e)}"}), 500

if __name__ == "__main__":
    app.run(host="0.0.0.0", port=8000)
