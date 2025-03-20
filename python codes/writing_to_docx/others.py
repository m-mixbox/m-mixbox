from turtle import width
from docx.shared import Inches
from docx import Document
from flask import Flask, request, send_file, jsonify
import os
import random
import string
from io import BytesIO
from docx import Document
from docx.shared import Inches
from pathlib import Path
import traceback


# Open the existing Word document
local_path = os.path.dirname(__file__)
image_dir = os.path.join(local_path,"images")
output_dir = os.path.join(local_path,"generated")
values_list = []
def generate_random_filename(length=7):
    return ''.join(random.choices(string.ascii_letters + string.digits, k=length))

# Access the first table (index 0)
def others_format(local_path,image_dir,output_dir,values_list,random_filename,output_path):
        
        template_path = os.path.join(os.path.dirname(__file__), "others.docx")
        print(f"Template path: {template_path}")

        if not os.path.exists(template_path):
            return jsonify({"error": "Template file 'others Format.docx' not found."}), 404

        # Open the template
        doc = Document(template_path)

        table = doc.tables[0]
        table_input_position = [4, 4, 4, 4, 4, 4, 2, 4]

        for i in range(len(table.rows)):
            #row = table.rows[i]
            #row_data = [cell.text.strip() for cell in row.cells]
            #print(len(row_data))
            #print(row_data)
            if i != 0:
                table.cell(i, table_input_position[i]).text = values_list[i]
            #elif i==5 :
                #table.cell(i, 5).text = ""
            #elif i==6 :
                #table.cell(i, 5).text = ""
            #elif i == 7:
                #table.cell(i, 5).text = values_list[5]
            #elif i ==8 :
                #table.cell(i, 5).text = values_list[6]
            #elif i==23:
                #table.cell(i, 5).text = ""

        #print("+++++++++++++++++++++++")
        #table_input_position_2 = [3,3,3,3,3,3,3,3,3,3,3,3,3,3,3,3,3]
        table_2 = doc.tables[1]
        #table_input_position_2 = [2, 2, 2]
        for i in range(len(table_2.rows)):
            #print(i)
            #row = table_2.rows[i]
            #row_data = [cell.text.strip() for cell in row.cells]
            #print(len(row_data))
            #print(row_data)
            if i !=0 or i !=9 or i != 15:
                table_2.cell(i, 3).text = values_list[i+9]
            #elif i == 16:
                #table_2.cell(i, 2).text = ""
            #elif i == 18:
                #table_2.cell(i, 2).text = ""
        #print("+++++++++++++++++++++++")
        #table_input_position_3 = [4,4,4,4,4,4,4,4,4]
        table_3 = doc.tables[2]
        #table_input_position_2 = [2, 2, 2]
        for i in range(len(table_3.rows)):
            #row = table_3.rows[i]
            #row_data = [cell.text.strip() for cell in row.cells]
            #print(len(row_data))
            #print(row_data)
            if i !=0 or i != 5:
                table_3.cell(i, 4).text = values_list[i+25]
        
        table_4 = doc.tables[3]
        #table_input_position_2 = [2, 2, 2]
        for i in range(len(table_4.rows)):
            #row = table_4.rows[i]
            #row_data = [cell.text.strip() for cell in row.cells]
            #print(len(row_data))
            #print(row_data)
            table_3.cell(i, 1).text = values_list[i+36]

        #table_4 = doc.tables[3]
        #for row in table_4.rows:
            #for cell in row.cells:
                #for paragraph in cell.paragraphs:
                    #for run in paragraph.runs:
                        #run.clear()

        image_index = 0
        start_row = 0
        start_col = 0
        #image_paths = [os.path.join(image_dir,"1E2KS.jpg"),os.path.join(image_dir,"1X0I8.jpg"),os.path.join(image_dir,"2A1VW.jpg"),os.path.join(image_dir,"2GW4H.jpg")]

        #for row_idx in range(start_row, len(table_4.rows)):
            #for col_idx in range(start_col, len(table_4.columns)):
                #if image_index >= len(image_paths):
                    #break
                #cell = table_4.cell(row_idx, col_idx)
                #paragraph = cell.paragraphs[0]
                #run = paragraph.add_run()
                #try:
                    #run.add_picture(image_paths[image_index], width=Inches(4.06))
                #except Exception as e:
                    #print(f"Error adding image {image_paths[image_index]}: {e}")
                #image_index += 1
            #if image_index >= len(image_paths):
                #break

        # Save the document to a file
        #doc.save(output_path)
        return output_path

        print(f"File saved to {output_path}")
random_filename = generate_random_filename()
output_path = os.path.join(output_dir,random_filename)
abc = others_format(local_path,image_dir,output_dir,values_list,random_filename,output_path)