import pytesseract
import shutil
import os
import random
import datetime
from datetime import date
import pytz
import docx
try:
    from PIL import Image
except ImportError:
    import Image

def create_directory():
        current_time = datetime.datetime.now(pytz.timezone('Asia/Kolkata'))
        time = str(current_time.hour) + '-' + str(current_time.minute) +'-' + str(current_time.second)
        directory = str(date.today())+ ' ' +time
        parent_dir = "D:/microsoft word/"
        path = os.path.join(parent_dir, directory) 
        error = ''
        try:
            os.makedirs(path, exist_ok = True)
            print("Directory '%s' created successfully" % directory) 
            error = ''
        except OSError as error:
            print(error)
        return error,path
path = []
for i in range(1):
     
    extractedInformation = pytesseract.image_to_string(Image.open(r"C:\Users\MBSPL-Ayush\Documents\Scan2024-02-19_174849_006.jpg"))
    with open('docx_file.txt', 'w') as f:
        f.write(extractedInformation)
#print(extractedInformation)
    doc = docx.Document()
    doc.add_paragraph().add_run(extractedInformation)
    error,path = create_directory()
    filename = path + '/'+'data'+'.docx'
    doc.save(filename)