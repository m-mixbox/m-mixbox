import pytesseract
import shutil
import os
import random
import datetime
from datetime import date as dt
import pytz
import docx
try:
    from PIL import Image
except ImportError:
    import Image

def create_directory():
        current_time = datetime.datetime.now(pytz.timezone('Asia/Kolkata'))
        time = str(current_time.hour) + '-' + str(current_time.minute) +'-' + str(current_time.second)
        directory = str(date.today())+ ' ' +time
        parent_dir = "D:/microsoft word/"
        path = os.path.join(parent_dir, directory) 
        error = ''
        try:
            os.makedirs(path, exist_ok = True)
            print("Directory '%s' created successfully" % directory) 
            error = ''
        except OSError as error:
            print(error)
        return error,path

def extract_text(file_path,save_file):
    extractedInformation = pytesseract.image_to_string(Image.open(file_path))
    current_time = datetime.datetime.now(pytz.timezone('Asia/Kolkata'))
    dir_name = str(dt.today())+ '____'+ str(current_time.hour) + '-' + str(current_time.minute) +'-' + str(current_time.second)
    dir_name.replace('-','_')
    #with open(save_file+'/'+'docx_file.txt', 'w') as f:
       # f.write(extractedInformation)
#print(extractedInformation)
    doc = docx.Document()
    doc.add_paragraph().add_run(extractedInformation)
    #print('file created')

    #error,path = create_directory()
    filename = save_file + '/'+dir_name+'.docx'
    doc.save(filename)
    return dir_name+'.docx'